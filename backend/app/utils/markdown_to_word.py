"""
Markdown转Word文档工具类
用于将Markdown内容转换为Word文档格式
"""
import io
import logging
from typing import Optional
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from markdown import markdown
from bs4 import BeautifulSoup, NavigableString, Tag

logger = logging.getLogger(__name__)


def markdown_to_word(markdown_content: str) -> bytes:
    """
    将Markdown内容转换为Word文档
    
    Args:
        markdown_content: Markdown格式的文本内容
        
    Returns:
        Word文档的二进制内容
    """
    try:
        # 1. Markdown -> HTML
        html = markdown(markdown_content, extensions=['extra', 'nl2br', 'tables'])
        
        # 2. 创建Word文档
        doc = Document()
        
        # 3. 解析HTML并转换为Word
        soup = BeautifulSoup(html, 'html.parser')
        _process_html_elements(doc, soup)
        
        # 4. 保存到字节流
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        
        return output.read()
    
    except Exception as e:
        logger.error(f"Markdown转Word失败: {str(e)}")
        raise


def _process_html_elements(doc: Document, soup: BeautifulSoup):
    """处理HTML元素并转换为Word格式"""
    for element in soup.children:
        if isinstance(element, NavigableString):
            text = str(element).strip()
            if text:
                doc.add_paragraph(text)
        elif isinstance(element, Tag):
            _process_tag(doc, element)


def _process_tag(doc: Document, tag: Tag, parent_paragraph=None):
    """处理HTML标签"""
    tag_name = tag.name
    
    if tag_name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
        # 标题
        level = int(tag_name[1])
        text = tag.get_text().strip()
        if text:
            para = doc.add_heading(text, level=level)
            
            # 设置标题样式
            if level == 1:
                para.runs[0].font.size = Pt(18)
                para.runs[0].font.color.rgb = RGBColor(37, 99, 235)  # 科技蓝色 #2563EB
                para.runs[0].font.bold = True
            elif level == 2:
                para.runs[0].font.size = Pt(16)
                para.runs[0].font.bold = True
            elif level == 3:
                para.runs[0].font.size = Pt(14)
                para.runs[0].font.bold = True
    
    elif tag_name == 'p':
        # 段落
        para = doc.add_paragraph()
        para.paragraph_format.line_spacing = 1.5  # 1.5倍行距
        _process_inline_elements(para, tag)
    
    elif tag_name in ['ul', 'ol']:
        # 列表
        for li in tag.find_all('li', recursive=False):
            text = li.get_text().strip()
            if text:
                para = doc.add_paragraph(text, style='List Bullet' if tag_name == 'ul' else 'List Number')
                para.paragraph_format.line_spacing = 1.3
    
    elif tag_name == 'blockquote':
        # 引用
        text = tag.get_text().strip()
        if text:
            para = doc.add_paragraph(text)
            para.paragraph_format.left_indent = Inches(0.5)
            para.paragraph_format.line_spacing = 1.3
            for run in para.runs:
                run.font.italic = True
                run.font.color.rgb = RGBColor(100, 116, 139)  # 灰色
    
    elif tag_name == 'pre':
        # 代码块
        code = tag.get_text().strip()
        if code:
            para = doc.add_paragraph(code)
            para.paragraph_format.left_indent = Inches(0.5)
            para.paragraph_format.line_spacing = 1.2
            for run in para.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0, 0, 0)
    
    elif tag_name == 'table':
        # 表格
        _process_table(doc, tag)
    
    elif tag_name == 'hr':
        # 分隔线
        doc.add_paragraph('─' * 50)
    
    else:
        # 其他标签，递归处理子元素
        for child in tag.children:
            if isinstance(child, Tag):
                _process_tag(doc, child)
            elif isinstance(child, NavigableString):
                text = str(child).strip()
                if text:
                    doc.add_paragraph(text)


def _process_inline_elements(paragraph, tag):
    """处理段落内的行内元素（粗体、斜体等）"""
    for element in tag.children:
        if isinstance(element, NavigableString):
            text = str(element)
            if text:
                run = paragraph.add_run(text)
                run.font.size = Pt(12)
        elif isinstance(element, Tag):
            text = element.get_text()
            if text:
                run = paragraph.add_run(text)
                run.font.size = Pt(12)
                
                # 处理样式
                if element.name in ['strong', 'b']:
                    run.font.bold = True
                elif element.name in ['em', 'i']:
                    run.font.italic = True
                elif element.name == 'code':
                    run.font.name = 'Courier New'
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(220, 38, 38)  # 红色
                elif element.name == 'a':
                    run.font.color.rgb = RGBColor(37, 99, 235)  # 蓝色
                    run.font.underline = True


def _process_table(doc: Document, table_tag: Tag):
    """处理HTML表格"""
    try:
        # 获取所有行
        rows = table_tag.find_all('tr')
        if not rows:
            return
        
        # 创建Word表格
        max_cols = max(len(row.find_all(['th', 'td'])) for row in rows)
        word_table = doc.add_table(rows=len(rows), cols=max_cols)
        word_table.style = 'Light Grid Accent 1'
        
        # 填充数据
        for row_idx, row in enumerate(rows):
            cells = row.find_all(['th', 'td'])
            for col_idx, cell in enumerate(cells):
                if col_idx < max_cols:
                    text = cell.get_text().strip()
                    word_table.rows[row_idx].cells[col_idx].text = text
                    
                    # 表头样式
                    if cell.name == 'th':
                        for paragraph in word_table.rows[row_idx].cells[col_idx].paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
    
    except Exception as e:
        logger.warning(f"处理表格失败: {str(e)}")

