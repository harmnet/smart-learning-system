"""
AI智能创作API端点
用于生成、保存和导出AI创作的教学内容
"""
from typing import Any, List, Optional
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from fastapi.responses import StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.future import select
from pydantic import BaseModel
import io
import os
import uuid
import logging
from datetime import datetime
from urllib.parse import quote

from app.db.session import get_db
from app.models.llm_config import LLMConfig
from app.models.knowledge_graph import KnowledgeGraph, KnowledgeNode
from app.models.teaching_resource import TeachingResource
from app.models.base import User
from app.utils.oss_client import oss_client
from app.utils.markdown_to_word import markdown_to_word
import httpx
import json

router = APIRouter()
logger = logging.getLogger(__name__)

# 文件大小限制
MAX_FILE_SIZE = 20 * 1024 * 1024  # 20MB
MAX_CONTENT_LENGTH = 60000  # 60k字符


class GenerateRequest(BaseModel):
    knowledge_point: str
    graph_id: int
    teacher_id: int


class GenerateResponse(BaseModel):
    success: bool
    content: Optional[str] = None
    error: Optional[str] = None


class SaveRequest(BaseModel):
    teacher_id: int
    resource_name: str
    markdown_content: str
    knowledge_point: str
    folder_id: Optional[int] = None


class SaveResponse(BaseModel):
    success: bool
    resource_id: Optional[int] = None
    error: Optional[str] = None


# ==================== 辅助函数 ====================

async def parse_file_content(file: UploadFile) -> str:
    """解析上传文件的内容"""
    try:
        content_bytes = await file.read()
        
        # 根据文件类型解析
        filename = file.filename.lower()
        
        if filename.endswith('.txt') or filename.endswith('.md'):
            # 文本文件
            try:
                return content_bytes.decode('utf-8')
            except UnicodeDecodeError:
                try:
                    return content_bytes.decode('gbk')
                except:
                    return "[文本文件编码错误]"
        
        elif filename.endswith('.pdf'):
            # PDF文件
            try:
                from PyPDF2 import PdfReader
                pdf_file = io.BytesIO(content_bytes)
                reader = PdfReader(pdf_file)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                return text
            except Exception as e:
                logger.error(f"PDF解析失败: {str(e)}")
                return f"[PDF文件解析失败: {str(e)}]"
        
        elif filename.endswith(('.doc', '.docx')):
            # Word文件
            try:
                from docx import Document
                doc_file = io.BytesIO(content_bytes)
                doc = Document(doc_file)
                text = ""
                for para in doc.paragraphs:
                    text += para.text + "\n"
                return text
            except Exception as e:
                logger.error(f"Word文件解析失败: {str(e)}")
                return f"[Word文件解析失败: {str(e)}]"
        
        else:
            return f"[不支持的文件类型: {filename}]"
    
    except Exception as e:
        logger.error(f"文件内容解析失败: {str(e)}")
        return f"[文件读取失败: {str(e)}]"


def build_knowledge_tree_text(nodes: List[dict], level: int = 0) -> str:
    """构建知识图谱树形文本结构"""
    result = ""
    for node in nodes:
        indent = "  " * level
        result += f"{indent}- {node['node_name']}\n"
        if node.get('children'):
            result += build_knowledge_tree_text(node['children'], level + 1)
    return result


async def call_openai_compatible(config: LLMConfig, prompt: str) -> str:
    """调用OpenAI兼容的API"""
    endpoint_url = config.endpoint_url
    
    # 检查endpoint_url是否已经包含chat/completions
    if not endpoint_url.endswith('/chat/completions'):
        endpoint_url = f"{endpoint_url.rstrip('/')}/chat/completions"
    
    headers = {
        "Authorization": f"Bearer {config.api_key}",
        "Content-Type": "application/json"
    }
    
    data = {
        "model": config.model_name,
        "messages": [
            {"role": "user", "content": prompt}
        ],
        "temperature": 0.7,
        "max_tokens": 4000
    }
    
    try:
        # 增加超时到5分钟，AI生成长文本需要较长时间
        async with httpx.AsyncClient(timeout=300.0) as client:
            logger.info(f"调用LLM API: {endpoint_url}, 模型: {config.model_name}")
            response = await client.post(endpoint_url, json=data, headers=headers)
            response.raise_for_status()
            result = response.json()
            logger.info(f"LLM API调用成功，返回内容长度: {len(result['choices'][0]['message']['content'])}")
            return result["choices"][0]["message"]["content"]
    except httpx.TimeoutException as e:
        error_msg = f"AI生成超时（超过5分钟），请稍后重试或简化提示词"
        logger.error(f"LLM API超时: {str(e)}")
        raise Exception(error_msg)
    except httpx.HTTPStatusError as e:
        error_msg = f"AI生成失败: HTTP {e.response.status_code}"
        logger.error(f"LLM API HTTP错误: {e.response.status_code} - {e.response.text[:200]}")
        raise Exception(error_msg)
    except KeyError as e:
        error_msg = f"AI返回数据格式错误"
        logger.error(f"LLM API返回数据解析失败: {str(e)}")
        raise Exception(error_msg)
    except Exception as e:
        logger.error(f"LLM API调用异常: {type(e).__name__} - {str(e)}")
        raise


async def call_aliyun_qwen(config: LLMConfig, prompt: str) -> str:
    """调用阿里云通义千问API"""
    import dashscope
    
    dashscope.api_key = config.api_key
    
    response = dashscope.Generation.call(
        model=config.model_name or "qwen-turbo",
        prompt=prompt,
        max_tokens=4000
    )
    
    if response.status_code == 200:
        return response.output.text
    else:
        raise Exception(f"通义千问API调用失败: {response.message}")


async def call_wenxin(config: LLMConfig, prompt: str) -> str:
    """调用百度文心一言API"""
    # 文心一言需要access_token
    token_url = "https://aip.baidubce.com/oauth/2.0/token"
    
    async with httpx.AsyncClient(timeout=30.0) as client:
        # 获取access_token
        token_params = {
            "grant_type": "client_credentials",
            "client_id": config.api_key,
            "client_secret": config.api_secret
        }
        token_response = await client.post(token_url, params=token_params)
        token_response.raise_for_status()
        access_token = token_response.json()["access_token"]
        
        # 调用生成API
        api_url = f"{config.endpoint_url}?access_token={access_token}"
        data = {
            "messages": [
                {"role": "user", "content": prompt}
            ]
        }
        response = await client.post(api_url, json=data)
        response.raise_for_status()
        result = response.json()
        return result["result"]


# ==================== API端点 ====================

@router.post("/generate", response_model=GenerateResponse)
async def ai_generate_content(
    knowledge_point: str = Form(...),
    graph_id: int = Form(...),
    teacher_id: int = Form(...),
    custom_prompt: str = Form(default=""),
    selected_resource_ids: str = Form(default=""),  # 逗号分隔的ID列表
    auxiliary_files: List[UploadFile] = File(default=[]),
    db: AsyncSession = Depends(get_db),
) -> Any:
    """
    AI生成教学内容
    """
    try:
        # 1. 获取激活的LLM配置
        result = await db.execute(
            select(LLMConfig).where(LLMConfig.is_active == True)
        )
        config = result.scalars().first()
        
        if not config:
            return GenerateResponse(
                success=False,
                error="未找到激活的LLM配置，请先在管理后台配置并启用LLM"
            )
        
        if not config.api_key:
            return GenerateResponse(
                success=False,
                error="LLM配置的API Key未设置"
            )
        
        # 2. 获取知识图谱树形结构
        graph_result = await db.execute(
            select(KnowledgeGraph).where(
                KnowledgeGraph.id == graph_id,
                KnowledgeGraph.teacher_id == teacher_id
            )
        )
        graph = graph_result.scalars().first()
        
        if not graph:
            return GenerateResponse(
                success=False,
                error="知识图谱不存在"
            )
        
        # 获取知识图谱的所有节点
        nodes_result = await db.execute(
            select(KnowledgeNode).where(
                KnowledgeNode.graph_id == graph_id
            ).order_by(KnowledgeNode.id)
        )
        all_nodes = nodes_result.scalars().all()
        
        # 构建树形结构
        def build_tree(parent_id: Optional[int]) -> List[dict]:
            children = []
            for node in all_nodes:
                if node.parent_id == parent_id:
                    node_dict = {
                        'id': node.id,
                        'node_name': node.node_name,
                        'children': build_tree(node.id)
                    }
                    children.append(node_dict)
            return children
        
        tree = build_tree(None)
        knowledge_tree_text = build_knowledge_tree_text(tree)
        
        # 3. 处理选中的资源
        auxiliary_content = ""
        total_size = 0
        
        # 3.1 处理选中的教学资源
        if selected_resource_ids:
            try:
                resource_ids = [int(id.strip()) for id in selected_resource_ids.split(',') if id.strip()]
                for resource_id in resource_ids:
                    # 获取资源信息
                    resource_result = await db.execute(
                        select(TeachingResource).where(TeachingResource.id == resource_id)
                    )
                    resource = resource_result.scalars().first()
                    
                    if resource and resource.file_path:
                        try:
                            from app.utils.oss_client import get_oss_client
                            oss_client_instance = get_oss_client()
                            
                            if oss_client_instance and oss_client_instance.enabled:
                                # 从OSS读取文件内容
                                file_content = oss_client_instance.bucket.get_object(resource.file_path)
                                content_bytes = file_content.read()
                                
                                # 根据资源类型解析内容
                                resource_text = ""
                                if resource.resource_type in ['txt', 'markdown']:
                                    try:
                                        resource_text = content_bytes.decode('utf-8')
                                    except:
                                        try:
                                            resource_text = content_bytes.decode('gbk')
                                        except:
                                            resource_text = "[文本内容无法解析]"
                                elif resource.resource_type == 'pdf':
                                    try:
                                        from PyPDF2 import PdfReader
                                        pdf_file = io.BytesIO(content_bytes)
                                        reader = PdfReader(pdf_file)
                                        for page in reader.pages:
                                            resource_text += page.extract_text() + "\n"
                                    except:
                                        resource_text = "[PDF内容无法解析]"
                                elif resource.resource_type == 'word':
                                    try:
                                        from docx import Document
                                        doc_file = io.BytesIO(content_bytes)
                                        doc = Document(doc_file)
                                        for para in doc.paragraphs:
                                            resource_text += para.text + "\n"
                                    except:
                                        resource_text = "[Word内容无法解析]"
                                
                                if resource_text:
                                    auxiliary_content += f"\n\n--- {resource.resource_name} ---\n{resource_text}\n"
                                    total_size += len(resource_text)
                        except Exception as e:
                            logger.warning(f"Failed to read resource {resource_id}: {str(e)}")
            except Exception as e:
                logger.error(f"Failed to process selected resources: {str(e)}")
        
        # 3.2 处理上传的辅助文件
        for file in auxiliary_files:
            # 检查文件大小
            file_content = await file.read()
            file_size = len(file_content)
            total_size += file_size
            
            if total_size > MAX_FILE_SIZE:
                return GenerateResponse(
                    success=False,
                    error=f"辅助资料总大小超过限制（{MAX_FILE_SIZE / 1024 / 1024}MB）"
                )
            
            # 重置文件指针
            await file.seek(0)
            
            # 解析文件内容
            content = await parse_file_content(file)
            auxiliary_content += f"\n\n--- {file.filename} ---\n{content}\n"
        
        # 检查总内容长度
        if len(auxiliary_content) > MAX_CONTENT_LENGTH:
            return GenerateResponse(
                success=False,
                error="辅助资料内容太大，无法处理（限制60k字符）"
            )
        
        # 4. 构建AI提示词
        auxiliary_section = ""
        if auxiliary_content.strip():
            auxiliary_section = f"\n\n## 辅助参考资料\n{auxiliary_content}"
        
        custom_section = ""
        if custom_prompt.strip():
            custom_section = f"\n\n## 用户补充要求\n{custom_prompt.strip()}"
        
        prompt = f"""请根据以下知识点生成详细的教学讲解内容。

## 知识图谱结构
{knowledge_tree_text}

## 当前知识点
{knowledge_point}
{auxiliary_section}

## 基本要求
1. 内容要详实、系统、易懂，适合教学使用
2. 使用Markdown格式输出
3. 包含标题、段落、列表等结构元素
4. 字数控制在3000-5000字
5. 内容要有逻辑性，层次分明
6. 如果有辅助资料，请结合资料内容进行讲解
{custom_section}

请直接输出Markdown格式的教学内容，不要包含其他说明文字。
"""
        
        # 5. 调用LLM API
        provider_key = config.provider_key
        
        try:
            if provider_key == "aliyun_qwen":
                content = await call_aliyun_qwen(config, prompt)
            elif provider_key in ["deepseek", "kimi", "volcengine_doubao", "siliconflow"]:
                content = await call_openai_compatible(config, prompt)
            elif provider_key == "wenxin":
                content = await call_wenxin(config, prompt)
            else:
                return GenerateResponse(
                    success=False,
                    error=f"不支持的LLM提供商: {provider_key}"
                )
            
            # 清理内容（去除可能的markdown代码块标记）
            content = content.strip()
            if content.startswith('```markdown'):
                content = content[11:].strip()
            elif content.startswith('```'):
                content = content[3:].strip()
            if content.endswith('```'):
                content = content[:-3].strip()
            
            return GenerateResponse(
                success=True,
                content=content
            )
        
        except Exception as e:
            logger.error(f"LLM API调用失败: {str(e)}")
            return GenerateResponse(
                success=False,
                error=f"AI生成失败: {str(e)}"
            )
    
    except Exception as e:
        logger.error(f"AI生成内容失败: {str(e)}")
        return GenerateResponse(
            success=False,
            error=f"生成失败: {str(e)}"
        )


@router.post("/save", response_model=SaveResponse)
async def save_ai_content(
    request: SaveRequest,
    db: AsyncSession = Depends(get_db),
) -> Any:
    """
    保存AI生成的内容到系统
    """
    try:
        # 1. Markdown转Word
        word_bytes = markdown_to_word(request.markdown_content)
        
        # 2. 生成文件名
        filename = f"{request.resource_name}.docx"
        file_key = f"ai-creation/{request.teacher_id}/{uuid.uuid4().hex}_{filename}"
        
        # 3. 上传到OSS
        if not oss_client.enabled:
            return SaveResponse(
                success=False,
                error="OSS服务未启用"
            )
        
        try:
            # 上传文件
            oss_client.bucket.put_object(file_key, word_bytes)
            file_path = file_key
            file_size = len(word_bytes)
        except Exception as e:
            logger.error(f"OSS上传失败: {str(e)}")
            return SaveResponse(
                success=False,
                error=f"文件上传失败: {str(e)}"
            )
        
        # 4. 保存到数据库
        # 获取教师信息
        user_result = await db.execute(
            select(User).where(User.id == request.teacher_id)
        )
        user = user_result.scalars().first()
        
        if not user:
            return SaveResponse(
                success=False,
                error="教师不存在"
            )
        
        # 创建教学资源记录
        resource = TeachingResource(
            teacher_id=request.teacher_id,
            resource_name=request.resource_name,
            original_filename=filename,
            file_path=file_path,
            file_size=file_size,
            resource_type="word",
            knowledge_point=request.knowledge_point,
            folder_id=request.folder_id,
            is_active=True
        )
        
        db.add(resource)
        await db.commit()
        await db.refresh(resource)
        
        return SaveResponse(
            success=True,
            resource_id=resource.id
        )
    
    except Exception as e:
        logger.error(f"保存AI内容失败: {str(e)}")
        await db.rollback()
        return SaveResponse(
            success=False,
            error=f"保存失败: {str(e)}"
        )


@router.post("/export")
async def export_ai_content(
    resource_name: str = Form(...),
    markdown_content: str = Form(...),
) -> Any:
    """
    导出AI生成的内容到本地
    """
    try:
        # 1. Markdown转Word
        word_bytes = markdown_to_word(markdown_content)
        
        # 2. 生成文件名
        filename = f"{resource_name}.docx"
        
        # 3. 返回文件流
        output = io.BytesIO(word_bytes)
        output.seek(0)
        
        return StreamingResponse(
            output,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename*=UTF-8''{quote(filename)}"
            }
        )
    
    except Exception as e:
        logger.error(f"导出AI内容失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"导出失败: {str(e)}")

